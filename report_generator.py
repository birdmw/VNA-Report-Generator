#!/usr/bin/python

# report_generator.py
# Matthew Bird
# 2/3/2016

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx.shared import Inches
import matlab.engine
import sys
import os
import pandas as pd
import docxtpl
import glob
import numpy as np
import seaborn as sns
sns.set_style("darkgrid")

# DEPENDENCIES:
#  Download and install Anaconda for Windows Python 2.7 - https://www.continuum.io/downloads#_windows
#  Download and install Matlab 2014b or later - https://www.mathworks.com/index.html?s_tid=gn_loc_drop
#   *After Matlab is installed, install MATLAB Engine for Python:
#    cd "matlabroot\extern\engines\python"
#    python setup.py install
#  Download and install docxtpl 0.1.9 or later - https://pypi.python.org/pypi/docxtpl/0.1.9
#   *this can typically be done by running cmd.exe as administrator and entering this command into the prompt:
#    C:\>pip install docxtpl

# USAGE:
#  Example:
#  C:\>python report_generator.py


class Generator(object):
    # Generator Class - This object combines to interact with several files:
    # (info.csv, template.docx, guide.csv, and many .sNp files)
    def __init__(self):
        '''
        declaration of important variables which house relevant data
        :return:
        '''
        self.template_path = None
        self.guide_path = None
        self.info_path = None
        self.sNp_dir = None
        self.figure_string = None
        self.figures = None

        self.df_info = pd.DataFrame()
        self.d_info = {}
        self.tpl = None
        self.df_guide = pd.DataFrame()
        self.sNp_paths = None
        self.sNp = []

        self.ts_manager = None  # self.ts_manager.db[0][1][0] ====> DB, TSfile 0, S21
        self.plots = None
        self.stop_frequency = None
        self.start_frequency = None
        self.step_frequency = None

    def read_from_paths(self):
        '''
        Reads various files into internal structure
        :return:
        '''
        self.read_guide()
        self.read_template()
        self.read_info()
        self.read_sNps()

    def read_template(self):
        '''
        Grabs template.docx and pulls it into memory
        :return:
        '''
        file_path = self.template_path + "\\template.docx"
        if os.path.isfile(file_path):
            print "template.docx found at", file_path
            self.tpl = docxtpl.DocxTemplate(file_path)
        else:
            print "template.docx not found"

    def read_guide(self):
        '''
        Reads in the guide.csv file found, if none is found, generates one
        :return:
        '''
        if self.guide_path:
            file_path = str(self.guide_path) + "\\guide.csv"
            if os.path.isfile(file_path):
                print "guide.csv found at", file_path
                self.df_guide = pd.read_csv(file_path)
                self.set_ts_port_maps()
            else:
                print "guide.csv not found, generating one instead"

                self.generate_guide()
                self.df_guide.to_csv('guide.csv')
                self.df_guide = pd.read_csv(file_path)
                self.set_ts_port_maps()

    def generate_guide(self):
        '''
        generates a guide.csv file from internal structure
        :return:
        '''
        self.get_sNp_paths()
        figures = self.go_figure()
        col_names = ['', 'title', 'magphase']
        for i_p in range(len(self.sNp_paths)):
            col_names.append('sNp_'+str(i_p))
        self.df_guide = pd.DataFrame(columns=col_names)  # put in column names
        s = ['paths', '', '']
        for p in self.sNp_paths:
            s.append(p)
        df_s = pd.DataFrame([s], columns=col_names) # put in row 1 with paths
        self.df_guide=self.df_guide.append(df_s)
        port_map = ['portmap', '', '']
        for _ in self.sNp_paths:
            port_map.append('13_24')
        df_portmap = pd.DataFrame([port_map], columns=col_names)
        self.df_guide = self.df_guide.append(df_portmap)
        for i_pl in range(len(figures)):
            s = ["plot_"+str(i_pl), "plot_"+str(i_pl), self.figures[i_pl][0]]
            for _ in self.sNp_paths:
                s.append("s"+self.figures[i_pl][1]+"_"+self.figures[i_pl][2])
            self.df_guide = self.df_guide.append(pd.DataFrame([s], columns=col_names))
        self.df_guide = self.df_guide.set_index(self.df_guide.columns[0])
        return self.df_guide

    def get_sNp_paths(self):
        '''
        combs the sNp directory for sNp files and pulls them in
        :return:
        '''
        if os.path.isfile(self.info_path + "\\guide.csv"):
            self.sNp_paths = self.df_guide.iloc[0, 3:].values.tolist()
            print self.df_guide.iloc[0, 3:].values.tolist()
        else:
            self.sNp_paths = glob.glob(self.sNp_dir+"/*.s*p")
        return self.sNp_paths

    def go_figure(self):
        '''
        Parses the list of strings that make up plot s-parameter instructions
        :return:
        '''
        split_figures = self.figure_string.split(",")
        self.figures = []
        for i_f in range(len(split_figures)):
            parsed = list()
            parsed.append(split_figures[i_f][-1])
            if len(split_figures[i_f]) == 3:
                sx = split_figures[i_f][1]
                sy = split_figures[i_f][2]
            else:
                split_list = split_figures[i_f].split("_")
                sx = "".join(ch for ch in split_list[0] if ch in "0123456789")
                sy = "".join(ch for ch in split_list[1] if ch in "0123456789")
            parsed.append(sx)
            parsed.append(sy)
            self.figures.append(parsed)
        print self.figures
        return self.figures

    def read_info(self):
        '''
        Reads in the info.csv file from the designated directory
        :return:
        '''
        file_path = self.info_path + "\\info.csv"
        if os.path.isfile(file_path):
            print "info.csv found at", file_path
            self.df_info = pd.read_csv(file_path)
            self.d_info = self.df_info.set_index('key').iloc[:, :1].T.to_dict('list')
            for k in self.d_info.iterkeys():
                self.d_info[k] = self.d_info[k][0]
            print self.d_info
        else:
            print "info.csv not found"

    def set_ts_port_maps(self):
        '''
        sets the internal port maps (self.ts_port_maps) to be 0 or 1 based on internal structure
        :return:
        '''
        self.ts_port_maps = self.df_guide.iloc[1, 3:]
        for p in range(len(self.ts_port_maps)):
            if any(self.ts_port_maps[p] == port for port in ["13_24", 0, "0"]):
                self.ts_port_maps[p] = 0
            elif any(self.ts_port_maps[p] == port for port in ["12_34", 1, "1"]):
                self.ts_port_maps[p] = 1
            else:
                self.ts_port_maps[p] = 2
        print "portmaps:",str(np.array(self.ts_port_maps))

    def read_sNps(self):
        '''
        creates a local TouchStone manager initialized with the sNp files designated
        :return:
        '''
        self.sNp_paths = self.get_sNp_paths()
        self.set_ts_port_maps()
        self.ts_manager = TouchstoneManager(self.sNp_paths, self.ts_port_maps)

    def write_from_data(self):
        '''
        Fills in the template and generates an output document with data from internal structure
        :return:
        '''
        self.gen_plots()
        if not(self.tpl):
            self.tpl = docxtpl.DocxTemplate()
        else:
            self.d_info['StartFrequency'] = str(round(self.start_frequency*1000))+"MHz"
            self.d_info['StopFrequency'] = str(round(self.stop_frequency))+"GHz"
            self.d_info['StepFrequency'] = str(round(self.step_frequency*1000))+"MHz"

            if not(self.df_info is None):
                index = self.df_info['key'].tolist()
                dut_start_index = index.index('<DUT>')
                dut_end_index = index.index('</DUT>')
                port_start_index = index.index('<PORT>')
                port_end_index = index.index('</PORT>')

            dut_table = self.df_info.fillna('').iloc[dut_start_index+1:dut_end_index].as_matrix()
            port_table = self.df_info.fillna('').iloc[port_start_index+1:port_end_index].as_matrix()

            sd1 = self.tpl.new_subdoc()
            sd2 = self.tpl.new_subdoc()
            sd3 = self.tpl.new_subdoc()
            sd4 = self.tpl.new_subdoc()

            table1 = sd1.add_table(rows=len(dut_table), cols=len(dut_table[0]))
            table1.style = 'TableGrid'
            for row in range(len(dut_table)):
                for col in range(len(dut_table[0])):
                    table1.cell(row, col).text = dut_table[row][col]

            table2 = sd2.add_table(rows=len(port_table), cols=len(port_table[0]))
            table2.style = 'TableGrid'
            for row in range(len(port_table)):
                for col in range(len(port_table[0])):
                    table2.cell(row, col).text = port_table[row][col]

            table3 = sd3.add_table(rows=len(self.sNp_paths)+1, cols=2)
            table3.style = 'TableGrid'
            table3.cell(0, 1).text = "Trace Name"
            table3.cell(0, 0).text = "sNp file"

            for row in range(len(self.sNp_paths)):
                table3.cell(row+1, 1).text = "TS"+str(row)
                table3.cell(row+1, 0).text = self.ts_manager.file_titles[row]

            for i in range(self.count_plots()):
                sd4.add_picture('img'+str(i)+'.png', width=Inches(6.5))

            self.d_info['DUTIdentification'] = sd1
            self.d_info['PortConfiguration'] = sd2
            self.d_info['TestSummaryLegend'] = sd3
            self.d_info['TestSummary'] = sd4

            self.tpl.render(self.d_info)
            print "rendered"

        self.tpl.save("generated_doc.docx")

    def count_plots(self):
        '''
        simple counts the number of plots from the internal structure
        :return:
        '''
        print "# of plots is:", str(len(self.df_guide.index)-2)
        return len(self.df_guide.index)-2

    def gen_plots(self):
        '''
        Fires up matlab, performs s2sdd transform, and creates images for each plot
        :return:
        '''
        plot_count = self.count_plots()
        sNp_count = len(self.sNp_paths)
        for p in range(plot_count):
            plt.clf()
            y_arr, labels = [], []
            self.plots = []
            XY_list = []
            for t in range(sNp_count):
                parameter = self.df_guide['sNp_'+str(t)].loc[2+p]
                magphase = self.df_guide['magphase'].loc[2+p]
                title = self.df_guide['title'].loc[2+p]
                if not(pd.isnull(parameter)):
                    # self.ts_manager.db[0][1][0] ====> DB, TSfile 0, S21
                    if len(parameter) == 3:
                        X = int(parameter[1])-1
                        Y = int(parameter[2])-1
                    else:
                        split_list = parameter.split("_")
                        X = int("".join(ch for ch in split_list[0] if ch in "0123456789"))-1
                        Y = int("".join(ch for ch in split_list[1] if ch in "0123456789"))-1
                    XY_list.append([X, Y])
                    labels.append("TS"+str(t)+" S"+str(X+1)+str(Y+1))
                    x = np.array(self.ts_manager.ghz[t][X][Y])
                    if magphase == 'm':
                        y = np.array(self.ts_manager.db[t][X][Y])
                        plt.ylabel('Magnitude (dB)')
                    else:
                        y = np.array(self.ts_manager.deg[t][X][Y])
                        plt.ylabel('Phase (deg)')
                    y_arr.append(y)
            plots_arr = []
            if all(XY == XY_list[0] for XY in XY_list):
                plt.title('S'+str(X+1)+str(Y+1))
            else:
                plt.title('Mixed Parameters')
            plt.xlabel('Frequency (GHz)')
            for ya in range(len(y_arr)):
                plots_arr.append(plt.plot(x, y_arr[ya], label=labels[ya]))
            self.stop_frequency = x[-1]
            self.start_frequency = x[1]
            self.step_frequency = (self.stop_frequency - self.start_frequency)/float(len(x))
            plt.legend(loc='lower right')
            plt.savefig("img"+str(p))
            self.plots.append(plt)


class TouchstoneManager(object):
    #  A housing place for touchstone data
    #  Access like this:
    #  tsm = TouchstronManager(["Filepath1","Filepath2",...], port_maps)
    #
    #  self.ts_manager.db[0][1][0] ====> DB, TSfile 0, S21

    def __init__(self, ts_paths, port_maps):
        '''
        Declaration of important variables to read touchstone data into for plotting
        :param ts_paths:
        :param port_maps:
        :return:
        '''
        self.pm = port_maps
        self.db, self.deg, self.ghz = [], [], []
        self.eng = matlab.engine.start_matlab()
        self.file_titles = None
        for ts_p in range(len(ts_paths)):
            dbX_params = []
            degX_params = []
            ghzX_params = []
            filename = str(ts_paths[ts_p])
            print "file =", filename
            self.eng.eval("touchstone_filename = '"+filename+"';", nargout=0)
            self.eng.eval("s_obj = sparameters(touchstone_filename);", nargout=0)
            self.eng.eval("ghz = s_obj.Frequencies./1e9;", nargout=0)
            self.eng.eval("s_raw = s2sdd(s_obj.Parameters,"+str(self.pm[ts_p]+1)+");", nargout=0)
            for X in range(1, 3):
                dbY_params = []
                degY_params = []
                ghzY_params = []
                for Y in range(1, 3):
                    self.eng.eval("s_plot = squeeze(s_raw("+str(X)+","+str(Y)+",:));", nargout=0)
                    self.eng.eval("y_db = 20 * log10(abs(s_plot));", nargout=0)
                    self.eng.eval("y_deg = (180./pi).*(angle(s_plot));", nargout=0)
                    db = np.array(self.eng.workspace['y_db']).T[0]
                    deg = np.array(self.eng.workspace['y_deg']).T[0]
                    ghz = np.array(self.eng.workspace['ghz']).T[0]
                    dbY_params.append(db)
                    degY_params.append(deg)
                    ghzY_params.append(ghz)
                dbX_params.append(dbY_params)
                degX_params.append(degY_params)
                ghzX_params.append(ghzY_params)
            self.db.append(dbX_params)
            self.deg.append(degX_params)
            self.ghz.append(ghzX_params)
        self.eng.quit()
        self.file_titles = self.get_sNp_titles(ts_paths)

    def earliest_difference(self, list_of_strings, reverse = False):
        '''
        Helper function that takes a list of strings and finds the first point they dont have in common
        :param list_of_strings:
        :param reverse:
        :return:
        '''
        min_length = len(min(list_of_strings, key=len))
        for i in range(min_length):
            if reverse == False:
                ch = list_of_strings[0][i]
            else:
                ch = list_of_strings[0][-i-1]
            for str in list_of_strings:
                if reverse == False:
                    if ch != str[i]:
                        return i
                else:
                    if ch != str[-i-1]:
                        return -i

    def get_sNp_titles(self, string_list):
        '''
        creates a list of distinct titles for each sNp file
        :param string_list:
        :return:
        '''
        first = self.earliest_difference(string_list, reverse=False)
        last = self.earliest_difference(string_list, reverse=True)
        names = [x[first:last] for x in string_list]
        return names


def main():
    '''
    The main function that is called them the program is invoked as main
    :return:
    '''
    g = Generator()
    if 'help' in sys.argv:
        print
        print "All parameters are optional, but the default path right now is", os.getcwd()
        print
        print "A short list of available commands (do not include filename in paths):"
        print "  i, info.csv path"
        print "  t, template.docx path"
        print "  g, guide.csv path"
        print "  d, directory with .sNp files"
        print "  p, plots to generate, ie: 'python report_generator.py p S1_1m,S1_2p,S2_2m'"
        print "     m is for magnitude, p is for phase"
        print
        print "USAGE EXAMPLE:"
        print "  C:\>python report_generator.py i c:\path\\to\info t c:\path\\to\\template"
        print
    else:
        #  Setting paths for input files from command line or uses defaults if no command is present
        g.info_path = sys.argv[sys.argv.index('i')+1] \
            if ('i' in sys.argv and sys.argv.index('i') != len(sys.argv)-1) else os.getcwd()
        g.template_path = sys.argv[sys.argv.index('t')+1] \
            if ('t' in sys.argv and sys.argv.index('t') != len(sys.argv)-1) else os.getcwd()
        g.guide_path = sys.argv[sys.argv.index('g')+1] \
            if ('g' in sys.argv and sys.argv.index('g') != len(sys.argv)-1) else os.getcwd()
        g.sNp_dir = sys.argv[sys.argv.index('d')+1] \
            if ('d' in sys.argv and sys.argv.index('d') != len(sys.argv)-1) else os.getcwd()
        g.figure_string = sys.argv[sys.argv.index('p')+1] \
            if ('p' in sys.argv and sys.argv.index('p') != len(sys.argv)-1) else "s1_1m,s2_2m,s1_2m,s2_1m"
        print g.figure_string
        g.read_from_paths()
        g.write_from_data()

if __name__ == "__main__":
    main()
