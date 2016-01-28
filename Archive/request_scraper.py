import matlab.engine
import numpy as np
from Tkinter import *
import tkFileDialog
import openpyxl
from functools import partial
import matplotlib.pyplot as plt
import string
import unicodedata
from docx import *
from docx.shared import *
from docxtpl import *

from docx.enum.style import WD_BUILTIN_STYLE
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
import seaborn as sns
from scipy import signal

import skrf as rf #<--- this would be nice but plotting does not work


class GUI(object):
    def __init__(self):
        self.root = Tk()

        #main variables
        self.request_filepath = StringVar()
        self.ts_count_var = IntVar()
        self.plot_count_var = IntVar()

        #request variables
        self.request_doc = None


        #touchstone variables
        self.ts_paths = []
        self.ts_entry_boxes = []
        self.ts_port_maps = []
        self.ts_transforms = []
        self.ts_manager = None #self.ts_manager.db[0][1][0] ====> DB, TSfile 0, S21

        #plot variables
        self.p_bode_domain = []
        self.plot_ts_choices = list(list()) #[plt][ts] ====> 0 or 1
        self.plot_param_choices = []
        self.plots = []

        #network properties
        self.start_frequency = None
        self.stop_frequency = None
        self.step_frequency = None
        self.IFBW = "1 KHz (PNA default)"
        self.averaging = "Off (N/A)"
        self.power_level = "-8 dbm"


    def askopenfile_rq(self):
        a = tkFileDialog.askopenfilename()
        self.request_filepath.set(str(a))
        if hasattr(self, 'window2'):
            if hasattr(self, 'window3'):
                self.window3.lift()
            else:
                self.window2.lift()

    def askopenfile_ts(self, ts_number):
        self.ts_paths[ts_number].set(str(tkFileDialog.askopenfilename()))
        if hasattr(self, 'window2'):
            if hasattr(self, 'window3'):
                self.window3.lift()
            else:
                self.window2.lift()

    def close_gui(self):
        self.root.destroy()
        self.gen_plots()
        self.gen_from_template()
        #self.gen_doc(title="output")
        #print str(self.request_filepath.get())

    def third_window(self):
        self.ts_manager = TouchstoneManager(self.ts_paths, self.ts_port_maps)
        self.window3 = Toplevel()
        Label(self.window3, text="Plot Options").grid(row=0, column=0)
        Label(self.window3, text="Mag").grid(row=0, column=1)
        Label(self.window3, text="Phase").grid(row=0, column=2)
        for ts_p_i in range(len(self.ts_paths)):
            Label(self.window3, text="TS" + str(ts_p_i)).grid(row=0, column=3+ts_p_i)
        for plt_i in range(int(self.plot_count_var.get())):
            Label(self.window3, text="Plot " + str(plt_i)).grid(row=plt_i+1, column=0)
        self.p_bode_domain = []
        self.plot_ts_choices = [[0]*int(self.ts_count_var.get()) for _ in xrange(int(self.plot_count_var.get()))]
        for i in range(len(self.plot_ts_choices)):
            for j in range(len(self.plot_ts_choices[i])):
                self.plot_ts_choices[i][j] = IntVar()
        for plt_i in range(int(self.plot_count_var.get())):
            self.p_bode_domain.append(IntVar())
            Radiobutton(self.window3, text="Mag", value=0, variable=self.p_bode_domain[plt_i]).grid(row=plt_i+1, column=1)
            Radiobutton(self.window3, text="Phase", value=1, variable=self.p_bode_domain[plt_i]).grid(row=plt_i+1, column=2)
            for ts_i in range(int(self.ts_count_var.get())):
                Checkbutton(self.window3, variable=self.plot_ts_choices[plt_i][ts_i]).grid(row=1+plt_i, column=3+ts_i)
                self.plot_ts_choices[plt_i][ts_i].set(1)
            parameters = ["S11", "S12", "S21", "S22"]
            self.plot_param_choices.append(IntVar())
            for p in range(len(parameters)):
                Radiobutton(self.window3, text=parameters[p], value=p, variable=self.plot_param_choices[-1]).grid(
                        row=plt_i+1, column=3+int(self.ts_count_var.get())+p)
        Button(self.window3, text='Ok', command=self.close_gui).grid(row=int(self.plot_count_var.get())+1,
                                                                column=int(self.ts_count_var.get()))
        self.window3.lift()

    def second_window(self):
        self.request_doc = RequestDoc(str(self.request_filepath.get()))
        self.window2 = Toplevel()
        Label(self.window2, text="Choose").grid(row=0, column=0)
        Label(self.window2, text="Touchstone Files").grid(row=0, column=2)
        Label(self.window2, text="Port Map").grid(row=0, column=3)
        #Label(self.window2, text="S2SDD").grid(row=0, column=5)
        self.window2_labels, ts_browse_buttons = [], []
        self.ts_paths, self.ts_entry_boxes, self.ts_port_maps, self.ts_transforms = [], [], [], []
        for i in xrange(int(self.ts_count_var.get())):
            self.ts_paths.append(StringVar())
            self.window2_labels.append(Label(self.window2, text="TS "+str(i)).grid(row=i+1, column=0))
            ts_browse_buttons.append(Button(self.window2, text='Browse...',
                                            command=partial(self.askopenfile_ts, i)).grid(row=i+1, column=1))
            self.ts_entry_boxes.append(Entry(self.window2, textvariable=self.ts_paths[i]).grid(row=i+1, column=2))
            self.ts_port_maps.append(IntVar())
            self.ts_transforms.append(IntVar())
            Radiobutton(self.window2, text="1,3 to 2,4", value=0, variable=self.ts_port_maps[i]).grid(row=i+1, column=3)
            Radiobutton(self.window2, text="1,2 to 3,4", value=1, variable=self.ts_port_maps[i]).grid(row=i+1, column=4)
            #Checkbutton(self.window2, variable=self.ts_transforms[i]).grid(row=i+1, column=5)
        Button(self.window2, text='Ok', command=self.third_window).grid(row=int(self.ts_count_var.get())+1, column=4)
        self.window2.lift()

    def first_window(self):
        Label(self.root, text="Request File: ").grid(row=0, column=0)
        Button(self.root, text=" Browse... ", command=self.askopenfile_rq).grid(row=0, column=1)
        Entry(self.root, textvariable=self.request_filepath).grid(row=0, column=2)
        Label(self.root, text="How many touchstone files? ").grid(row=1, column=1, sticky=E)
        Entry(self.root, textvariable=self.ts_count_var).grid(row=1, column=2)
        Label(self.root, text="How many plots? ").grid(row=2, column=1, sticky=E)
        Entry(self.root, textvariable=self.plot_count_var).grid(row=2, column=2)
        Button(self.root, text="Ok", command=self.second_window).grid(row=3, column=2)

        self.request_filepath.set('C:/Users/Administrator/PycharmProjects/VNAReportGenreator/A0_B0/request.xlsx')
        self.ts_count_var.set(1)
        self.plot_count_var.set(1)
        mainloop()

    def gen_plots(self):
        for p in range(self.plot_count_var.get()):
            plt.clf()
            y_arr, labels = [], []
            for t in range(self.ts_count_var.get()):
                if self.plot_ts_choices[p][t].get():
                    labels.append("TS"+str(t))
                    #self.ts_manager.db[0][1][0] ====> DB, TSfile 0, S21
                    X = 0 if self.plot_param_choices[p].get() < 2 else 1
                    Y = 0 if self.plot_param_choices[p].get() % 2 == 0 else 1
                    x = np.array(self.ts_manager.ghz[t][X][Y])
                    if self.p_bode_domain[p].get():
                        y = np.array(self.ts_manager.deg[t][X][Y])
                        plt.ylabel('Phase (deg)')
                    else:
                        y = np.array(self.ts_manager.db[t][X][Y])
                        plt.ylabel('Magnitude (dB)')
                    plt.title('S'+str(X+1)+str(Y+1))
                    y_arr.append(y)
            plots_arr = []
            plt.xlabel('Frequency (GHz)')
            for ya in range(len(y_arr)):
                plots_arr.append(plt.plot(x, y_arr[ya], label=labels[ya]))
            self.stop_frequency = x[-1]
            self.start_frequency = x[1]
            self.step_frequency = (self.stop_frequency - self.start_frequency)/float(len(x))
            plt.legend(loc='lower right')
            plt.savefig("img"+str(p))
            self.plots.append(plt)

    def gen_doc(self, title="Final_Report"):
        document = Document("header.docx")
        styles = document.styles

        #for img_num in range(len(self.plots)):

        paragraph = []
        pagebreaks=[]
        table = []
        paragraph.append(document.add_paragraph(text="Microsoft Signal Integrity Laboratory "
                                                "\n15101 NE 40th ST (Studio B) \nRedmond, "
                                                "WA 98052", style='Normal'))
        paragraph.append(document.add_paragraph(text="\n\n\nVNA TEST REPORT\n\n\n\n\n", style='Title'))

        for i_t in range(2):
            t = self.request_doc.tables[i_t]
            row_count, col_count = len(t), len(t[0])
            table.append(document.add_table(rows=row_count, cols=col_count, style='TableGrid'))
            for r in range(row_count):
                for c in range(col_count):
                    if str(t[r][c]) and str(t[r][c]) != 'None':
                        table[-1].cell(r, c).text = str(t[r][c])
                    else:
                        table[-1].cell(r, c).text = ''
            paragraph.append(document.add_paragraph(text="\n", style='Normal'))

        pagebreaks.append(document.add_page_break())
        paragraph.append(document.add_paragraph(text="Table Of Contents", style='Title'))
        paragraph.append(document.add_paragraph(text="1	Test Overview\n2	Test Setup\n   2.1	PNA Test Setup"
                                                     "\n   2.2	Hardware Apparatus\n   2.3	Software Apparatus\n   2.4	"
                                                     "DUT Identification\n   2.5	Port Configuration & Probe "
                                                     "Identification\n3   Test Summary"   ))
        pagebreaks.append(document.add_page_break())
        paragraph.append(document.add_paragraph(text="1 Test Overview"))
        paragraph[-1].bold = True
        paragraph[-1].size = Pt(22)

        document.save(title+".docx")

    def gen_from_template(self):
        doc = DocxTemplate("template.docx")
        rd = self.request_doc.table_dict
        context = {}
        for rd_td in rd.iterkeys():
            context[normalizeit(rd_td)] = rd[rd_td]
        print self.request_doc.tables[2][0][1]
        context['TestOverview'] = self.request_doc.tables[2][0][1]
        context['StartFrequency'] = str(round(self.start_frequency*1000))+"MHz"
        context['StopFrequency'] = str(round(self.stop_frequency))+"GHz"
        context['FrequencyStep'] = str(round(self.step_frequency*1000))+"KHz"
        for key in context.iterkeys():
            print key
        sd = doc.new_subdoc()
        t1_r, t1_c = len(self.request_doc.tables[3])-1, len(self.request_doc.tables[3][0])
        table = sd.add_table(rows=t1_r, cols=t1_c, style = 'TableGrid')
        for i in range(t1_r):
            for j in range(t1_c):
                txt=str(self.request_doc.tables[3][i][j])
                if not(txt) or txt=="None":
                    txt=''
                table.cell(i, j).text = txt

        context['DUTIdentification'] = sd

        doc.render(context)
        doc.save("generated_doc.docx")

    def run(self):
        self.first_window()


class RequestDoc(object):
    def __init__(self, request_path):
        self.req_path = request_path
        self.wb = openpyxl.load_workbook(self.req_path)
        self.ws = self.wb.active
        self.sheet_names = self.wb.get_sheet_names()
        self.boundaries = []
        self.tables = []
        self.table_dict = {}
        self.scrape_table_boundaries()
        self.scrape_tables()

    def scrape_table_boundaries(self):
        self.boundaries = []
        prev_b = 1
        for row in self.ws.iter_rows('B1:B100'):
            for cell in row:
                if cell.value and "end>" in cell.value:
                    self.boundaries.append((prev_b, cell.row))
                    prev_b = cell.row+1

    def scrape_tables(self):
        last_line = max(self.boundaries[-1])
        for i in range(1, last_line):
            key = normalizeit(str(self.ws['B'+str(i)].value))
            value = normalizeit(str(self.ws['C'+str(i)].value))
            self.table_dict[key] = value
        for i_b in range(len(self.boundaries)):
            if i_b in [2, 4]:
                if i_b == 4:
                    self.tables.append(self.parse_port_config(self.ws, self.boundaries[i_b][0],
                                                         self.boundaries[i_b][1]))
                else:
                    table_line = []
                    for rng in range(self.boundaries[i_b][0], self.boundaries[i_b][1]):
                        if "Test Overview" in self.ws['B'+str(rng)].value:
                            table_line.append((str(self.ws['B'+str(rng)].value), str(self.ws['B'+str(rng+1)].value)))
                    self.tables.append(table_line)
            else:
                table_line = []
                for rng in range(self.boundaries[i_b][0], self.boundaries[i_b][1]):
                    if self.ws['B'+str(rng)].value:
                        table_line.append((str(self.ws['B'+str(rng)].value), str(self.ws['C'+str(rng)].value)))
                self.tables.append(table_line)

    def parse_port_config(self, ws, north, south):
        #Special function to read in the matrix from request document under the Port Configuration section
        west = 100
        east = 1
        return_table = []
        for ro in range(north, south):
            for co in range(1,100):
                if not(ws.cell(row=ro, column=co).value is None):
                    east = max(east, co)
                    west = min(west, co)
        for ro in range(north, south):
            table_line = []
            for co in range(west, east):
                table_line.append(str(ws.cell(row=ro, column=co).value))
            return_table.append(table_line)
        return return_table

def normalizeit(myString):
    a = myString
    a = ''.join(ch for ch in a if ch in "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ./()[];',0123456789\\")
    a = a[:20]
    return a

class TouchstoneManager(object):
    ''' TouchstoneManager
      Access like this:
       tsm = TouchstronManager(["Filepath1","Filepath2",...], port_maps)
    '''
    def __init__(self, ts_paths, port_maps):
        self.pm = port_maps
        self.db, self.deg, self.ghz = [], [], []
        self.eng = matlab.engine.start_matlab()
        for ts_p in range(len(ts_paths)):
            dbX_params = []
            degX_params = []
            ghzX_params = []
            filename = str(ts_paths[ts_p].get())
            print "file =",filename
            self.eng.eval("touchstone_filename = '"+filename+"';", nargout=0)
            self.eng.eval("s_obj = sparameters(touchstone_filename);", nargout=0)
            self.eng.eval("ghz = s_obj.Frequencies./1e9;", nargout=0)
            self.eng.eval("s_raw = s2sdd(s_obj.Parameters,"+str(self.pm[ts_p].get()+1)+");", nargout=0)
            for X in range(1, 3):
                dbY_params = []
                degY_params = []
                ghzY_params = []
                for Y in range(1, 3):
                    self.eng.eval("s_plot = squeeze(s_raw("+str(X)+","+str(Y)+",:));", nargout=0)
                    self.eng.eval("y_db = 20 * log10(abs(s_plot));", nargout=0)
                    self.eng.eval("y_deg = (180./pi).*(angle(s_plot));", nargout=0)
                    db = np.array(self.eng.workspace['y_db']).T[0]
                    deg = np.array(self.eng.workspace['y_deg']).T[0]
                    ghz = np.array(self.eng.workspace['ghz']).T[0]
                    dbY_params.append(db)
                    degY_params.append(deg)
                    ghzY_params.append(ghz)
                dbX_params.append(dbY_params)
                degX_params.append(degY_params)
                ghzX_params.append(ghzY_params)
            self.db.append(dbX_params)
            self.deg.append(degX_params)
            self.ghz.append(ghzX_params)
        self.eng.quit()
        return

s2pfile = 'C:/Users/Administrator/Desktop/DATA_CHANNELA_DQ0_AB1.s2p'
s4pfile = 'C:/Users/Administrator/PycharmProjects/VNAReportGenreator/A0_B0/SI-489-1_SURFACE_PCIE0_LANE1_C1_P1-2-3_Sdd1-2.s4p'

if __name__ == "__main__":
    g = GUI()
    g.run()