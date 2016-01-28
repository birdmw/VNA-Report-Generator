import cPickle
from request_scraper import GUI
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

g = cPickle.load(open("rq_obj.p"))

document = Document()
styles = document.styles
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]

paragraph = document.add_paragraph("Hello DoCx")
paragraph.style = document.styles['Heading 1']

#document.add_picture(self.images[0])

document.save('output.docx')